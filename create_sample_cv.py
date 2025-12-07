#!/usr/bin/env python3
"""
Create a sample CV file in DOCX format for testing
"""

try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx library not found. Installing...")
    import subprocess
    import sys
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
        from docx.shared import Pt
        HAS_DOCX = True
        print("✅ python-docx installed successfully!")
    except:
        print("❌ Could not install python-docx. Please install manually:")
        print("   pip install python-docx")
        HAS_DOCX = False

if HAS_DOCX:
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('JOHN DOE - SOFTWARE ENGINEER', 0)
    title.alignment = 1  # Center alignment
    
    # Add contact information
    doc.add_heading('CONTACT INFORMATION', level=1)
    doc.add_paragraph('Email: john.doe@email.com')
    doc.add_paragraph('Phone: +1 (555) 123-4567')
    doc.add_paragraph('Location: San Francisco, CA')
    doc.add_paragraph('LinkedIn: linkedin.com/in/johndoe')
    
    # Add professional summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    doc.add_paragraph(
        'Experienced Software Engineer with 5+ years of expertise in full-stack development, '
        'specializing in Java, Kotlin, and Android application development. Proven track record '
        'of delivering high-quality mobile applications and web services.'
    )
    
    # Add technical skills
    doc.add_heading('TECHNICAL SKILLS', level=1)
    skills = [
        'Programming Languages: Java, Kotlin, Python, JavaScript, TypeScript',
        'Mobile Development: Android SDK, Android Studio, Material Design',
        'Backend: Spring Boot, Node.js, RESTful APIs, GraphQL',
        'Databases: PostgreSQL, MySQL, MongoDB',
        'Cloud Services: AWS, Firebase, Docker',
        'Version Control: Git, GitHub, GitLab',
        'Testing: JUnit, Espresso, Mockito'
    ]
    for skill in skills:
        p = doc.add_paragraph(skill, style='List Bullet')
    
    # Add professional experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    
    doc.add_heading('Senior Software Engineer | Tech Company Inc. | 2021 - Present', level=2)
    exp1 = [
        'Developed and maintained Android applications with 100K+ downloads',
        'Designed and implemented RESTful APIs using Spring Boot',
        'Collaborated with cross-functional teams to deliver features on time',
        'Reduced app crash rate by 40% through improved error handling',
        'Mentored junior developers and conducted code reviews'
    ]
    for item in exp1:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Software Engineer | StartupXYZ | 2019 - 2021', level=2)
    exp2 = [
        'Built native Android apps using Kotlin and Java',
        'Integrated third-party APIs and SDKs',
        'Implemented real-time features using WebSockets',
        'Participated in agile development processes'
    ]
    for item in exp2:
        doc.add_paragraph(item, style='List Bullet')
    
    # Add education
    doc.add_heading('EDUCATION', level=1)
    doc.add_heading('Bachelor of Science in Computer Science', level=2)
    doc.add_paragraph('University of California, Berkeley | 2015 - 2019')
    doc.add_paragraph('GPA: 3.8/4.0', style='List Bullet')
    doc.add_paragraph('Relevant Coursework: Data Structures, Algorithms, Software Engineering, Database Systems', style='List Bullet')
    
    # Add projects
    doc.add_heading('PROJECTS', level=1)
    
    doc.add_heading('Interview Practice App (2024)', level=2)
    proj1 = [
        'Developed an AI-powered interview practice application',
        'Integrated OpenAI API for real-time interview questions',
        'Implemented WebSocket communication for live sessions',
        'Technologies: Kotlin, Android SDK, Retrofit, Ktor'
    ]
    for item in proj1:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('E-Commerce Mobile App (2023)', level=2)
    proj2 = [
        'Built a full-stack e-commerce application',
        'Features: Product catalog, shopping cart, payment integration',
        'Technologies: Android, Spring Boot, PostgreSQL'
    ]
    for item in proj2:
        doc.add_paragraph(item, style='List Bullet')
    
    # Add certifications
    doc.add_heading('CERTIFICATIONS', level=1)
    doc.add_paragraph('Google Associate Android Developer (2022)', style='List Bullet')
    doc.add_paragraph('AWS Certified Developer - Associate (2021)', style='List Bullet')
    
    # Add languages
    doc.add_heading('LANGUAGES', level=1)
    doc.add_paragraph('English (Native)', style='List Bullet')
    doc.add_paragraph('Spanish (Conversational)', style='List Bullet')
    
    # Save the document
    output_path = r'C:\Users\Admin\Desktop\razzan\razzan\sample_cv.docx'
    doc.save(output_path)
    print(f"✅ Sample CV created successfully: {output_path}")
    print(f"📄 File size: {len(open(output_path, 'rb').read())} bytes")
    
else:
    print("\n❌ Cannot create DOCX file without python-docx library.")
    print("\n📋 Alternative Options:")
    print("1. Install python-docx: pip install python-docx")
    print("2. Use your own PDF/DOCX file and push it to emulator")
    print("3. Use Android Studio Device File Explorer to upload a file")

